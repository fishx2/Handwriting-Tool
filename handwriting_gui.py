import tkinter as tk
from tkinter import ttk, scrolledtext, filedialog, messagebox, colorchooser
import os
import threading
from PIL import Image, ImageTk
import io
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.lib.utils import ImageReader
from docx import Document
from docx.shared import Inches
import json
import datetime

# Thêm import cho kéo-thả
try:
    from tkinterdnd2 import DND_FILES, TkinterDnD
    DND_AVAILABLE = True
except ImportError:
    DND_AVAILABLE = False

# Import HandwritingGenerator từ file trước
import os
import random
import numpy as np
from PIL import Image, ImageDraw, ImageFont, ImageFilter
import textwrap
import requests

# Thêm dict ngôn ngữ mẫu
LANGS = {
    'vi': {
        'title': '🖋️ Handwriting Generator - Chuyển đổi văn bản thành chữ viết tay',
        'font_size': 'Kích thước chữ:',
        'choose_font': 'Chọn font:',
        'font_demo': 'Demo font:',
        'line_spacing': 'Khoảng cách dòng:',
        'margin': 'Lề trang:',
        'ink_color': 'Màu mực:',
        'paper_style': 'Kiểu giấy:',
        'paper_size': 'Kích thước giấy:',
        'create': '🎨 Tạo chữ viết tay',
        'save_settings': '💾 Lưu cài đặt',
        'open_file': '📁 Mở file',
        'export_png': '💾 PNG',
        'export_pdf': '📄 PDF',
        'export_word': '📝 Word',
        'input_text': '✍️ Văn bản đầu vào',
        'char_count': 'Số ký tự:',
        'preview': '👁️ Xem trước',
        'ready': 'Sẵn sàng',
        'generating': 'Đang tạo chữ viết tay...'
    },
    'en': {
        'title': '🖋️ Handwriting Generator - Convert text to handwriting',
        'font_size': 'Font size:',
        'choose_font': 'Choose font:',
        'font_demo': 'Font demo:',
        'line_spacing': 'Line spacing:',
        'margin': 'Page margin:',
        'ink_color': 'Ink color:',
        'paper_style': 'Paper style:',
        'paper_size': 'Paper size:',
        'create': '🎨 Generate handwriting',
        'save_settings': '💾 Save settings',
        'open_file': '📁 Open file',
        'export_png': '💾 PNG',
        'export_pdf': '📄 PDF',
        'export_word': '📝 Word',
        'input_text': '✍️ Input text',
        'char_count': 'Characters:',
        'preview': '👁️ Preview',
        'ready': 'Ready',
        'generating': 'Generating handwriting...'
    }
}

FONT_LANGS = [
    ('english', 'English'),
    ('vietnamese', 'Tiếng Việt'),
    ('japanese', '日本語'),
    ('chinese', '中文'),
    ('korean', '한국어'),
    ('hindi', 'हिन्दी'),
    ('thai', 'ภาษาไทย'),
    ('russian', 'Русский'),
    ('arabic', 'العربية'),
    ('french', 'Français'),
    ('german', 'Deutsch'),
    ('spanish', 'Español'),
    ('italian', 'Italiano'),
    ('greek', 'Ελληνικά'),
    ('hebrew', 'עברית'),
    ('turkish', 'Türkçe'),
    ('indonesian', 'Bahasa Indonesia'),
    ('malay', 'Bahasa Melayu'),
    ('lao', 'ລາວ'),
    ('khmer', 'ខ្មែរ'),
    ('bengali', 'বাংলা'),
    ('tamil', 'தமிழ்'),
    ('telugu', 'తెలుగు'),
    ('gujarati', 'ગુજરાતી'),
    ('kannada', 'ಕನ್ನಡ'),
    ('malayalam', 'മലയാളം'),
    ('punjabi', 'ਪੰਜਾਬੀ'),
    ('urdu', 'اردو'),
    ('burmese', 'မြန်မာ'),
    ('mongolian', 'Монгол'),
    ('nepali', 'नेपाली'),
    ('sinhala', 'සිංහල'),
    ('georgian', 'ქართული'),
    ('armenian', 'Հայերեն'),
]

LANG_LABELS = {
    'english': 'English',
    'vietnamese': 'Vietnamese/Tiếng Việt',
    'japanese': 'Japanese/日本語',
    'chinese': 'Chinese/中文',
    'korean': 'Korean/한국어',
    'hindi': 'Hindi/हिन्दी',
    'thai': 'Thai/ภาษาไทย',
    'russian': 'Russian/Русский',
    'arabic': 'Arabic/العربية',
    'french': 'French/Français',
    'german': 'German/Deutsch',
    'spanish': 'Spanish/Español',
    'italian': 'Italian/Italiano',
    'greek': 'Greek/Ελληνικά',
    'hebrew': 'Hebrew/עברית',
    'turkish': 'Turkish/Türkçe',
    'indonesian': 'Indonesian/Bahasa Indonesia',
    'malay': 'Malay/Bahasa Melayu',
    'lao': 'Lao/ລາວ',
    'khmer': 'Khmer/ខ្មែរ',
    'bengali': 'Bengali/বাংলা',
    'tamil': 'Tamil/தமிழ்',
    'telugu': 'Telugu/తెలుగు',
    'gujarati': 'Gujarati/ગુજરાતી',
    'kannada': 'Kannada/ಕನ್ನಡ',
    'malayalam': 'Malayalam/മലയാളം',
    'punjabi': 'Punjabi/ਪੰਜਾਬੀ',
    'urdu': 'Urdu/اردو',
    'burmese': 'Burmese/မြန်မာ',
    'mongolian': 'Mongolian/Монгол',
    'nepali': 'Nepali/नेपाली',
    'sinhala': 'Sinhala/සිංහල',
    'georgian': 'Georgian/ქართული',
    'armenian': 'Armenian/Հայերեն',
}

class HandwritingGenerator:
    def __init__(self, font_lang='english'):
        self.font_lang = font_lang
        self.fonts = []
        self.font_names = {}
        try:
            self.fallback_font_path = "C:/Windows/Fonts/seguiemj.ttf"
            ImageFont.truetype(self.fallback_font_path)
            print("INFO: Tải font dự phòng 'Segoe UI Emoji' thành công cho các ký tự đặc biệt.")
        except IOError:
            try:
                self.fallback_font_path = "C:/Windows/Fonts/arial.ttf"
                ImageFont.truetype(self.fallback_font_path)
                print("INFO: Tải font dự phòng 'Arial' thành công.")
            except IOError:
                self.fallback_font_path = None
                print("CẢNH BÁO: Không tìm thấy font dự phòng (Arial, Segoe UI). Ký tự đặc biệt có thể không hiển thị.")
        self.load_fonts()
        
    def load_fonts(self, font_lang=None):
        if font_lang:
            self.font_lang = font_lang
        # Đường dẫn thư mục font theo ngôn ngữ
        font_dir = os.path.join('fonts', self.font_lang)
        if not os.path.exists('fonts'):
            os.makedirs('fonts')
        if not os.path.exists(font_dir):
            os.makedirs(font_dir)
            # Ghi file hướng dẫn
            with open(os.path.join(font_dir, 'HUONG_DAN.txt'), 'w', encoding='utf-8') as f:
                f.write(f"Vui lòng sao chép các file font chữ .ttf hoặc .otf của bạn vào thư mục này cho ngôn ngữ {self.font_lang}.\n")
                f.write("Chỉ những font hỗ trợ đầy đủ ký tự của ngôn ngữ này mới được tải.\n")
        self.fonts = []
        self.font_names = {}
        font_files = [f for f in os.listdir(font_dir) if f.lower().endswith(('.ttf', '.otf'))]
        if not font_files:
            print(f"CẢNH BÁO: Không tìm thấy file font nào trong thư mục '{font_dir}'.")
        else:
            print(f"Bắt đầu quét và kiểm tra font chữ cho {self.font_lang}...")
            test_text = "abc" if self.font_lang == 'english' else (
                "ăâđêôơưàáảãạĐ" if self.font_lang == 'vietnamese' else (
                "日本語" if self.font_lang == 'japanese' else (
                "汉字" if self.font_lang == 'chinese' else (
                "한글" if self.font_lang == 'korean' else (
                "अआइई" if self.font_lang == 'hindi' else "abc")))))
            for font_file in font_files:
                font_path = os.path.join(font_dir, font_file)
                try:
                    test_font = ImageFont.truetype(font_path, 18, encoding='unic')
                    bbox = test_font.getbbox(test_text)
                    if bbox[2] <= 0 or bbox[3] <= 0:
                        raise ValueError("Font không vẽ được bounding box cho ký tự test.")
                    self.fonts.append(font_path)
                    display_name = os.path.splitext(font_file)[0]
                    self.font_names[font_path] = display_name
                    print(f"  [OK] Font '{display_name}' hợp lệ và đã được tải.")
                except Exception as e:
                    print(f"  [LỖI] Font '{font_file}' không hợp lệ. Bỏ qua. ({e})")
        if not self.fonts:
            print(f"CẢNH BÁO: Không có font hợp lệ nào được tìm thấy cho {self.font_lang}. Sẽ sử dụng font mặc định.")
            self.fonts.append(None)
            self.font_names[None] = 'Font mặc định (Lỗi)'
        
    def has_glyph(self, font, char):
        """Kiểm tra font có hỗ trợ vẽ một ký tự cụ thể hay không."""
        try:
            # getmask là phương pháp đáng tin cậy để kiểm tra sự tồn tại của glyph
            return font.getmask(char).getbbox() is not None
        except Exception:
            return False

    def add_natural_variations(self, draw, text, x, y, font, ink_color):
        """Thêm các biến thể tự nhiên cho chữ viết tay (cải tiến chuyên nghiệp)"""
        if not text.strip():
            return
        
        # Tạo hiệu ứng rung tay tự nhiên (tăng độ ngẫu nhiên)
        hand_tremor = random.randint(0, 4)  # Độ rung tay (0-4 pixel)
        
        # Thêm biến thể cho từng ký tự
        current_x = x
        for char in text:
            # Biến thể vị trí cho từng ký tự (tăng độ ngẫu nhiên)
            char_x = current_x + random.randint(-hand_tremor, hand_tremor)
            char_y = y + random.randint(-hand_tremor, hand_tremor)
            
            # Biến thể kích thước font nhẹ (tăng độ ngẫu nhiên)
            size_variation = random.randint(-3, 3)
            char_font_size = font.size + size_variation
            if char_font_size < 8:  # Đảm bảo font size tối thiểu
                char_font_size = 8
            
            try:
                # Tạo font với kích thước biến thể
                char_font = ImageFont.truetype(font.path, char_font_size, encoding='unic')
            except:
                # Fallback nếu không tạo được font mới
                char_font = font
            
            # Biến thể màu mực (đậm nhạt) - tăng độ ngẫu nhiên
            ink_variation = random.randint(-30, 30)
            char_ink_color = tuple(max(0, min(255, c + ink_variation)) for c in ink_color)
            
            # Hiệu ứng nét bút đứt đoạn (10% khả năng)
            if random.random() < 0.1:
                # Làm mờ màu để tạo hiệu ứng hết mực
                char_ink_color = tuple(max(0, min(255, c - 50)) for c in char_ink_color)
            
            # Vẽ ký tự với biến thể
            draw.text((char_x, char_y), char, font=char_font, fill=char_ink_color)
            
            # Tính toán khoảng cách đến ký tự tiếp theo (tăng độ ngẫu nhiên)
            bbox = draw.textbbox((0, 0), char, font=char_font)
            char_width = bbox[2] - bbox[0]
            
            # Thêm khoảng cách tự nhiên giữa các ký tự (tăng độ ngẫu nhiên)
            spacing_variation = random.randint(-3, 4)
            current_x += char_width + spacing_variation

    def add_ink_effects(self, image):
        """Thêm hiệu ứng mực và texture"""
        img_array = np.array(image)
        noise = np.random.normal(0, 5, img_array.shape)
        img_array = np.clip(img_array + noise, 0, 255).astype(np.uint8)
        result = Image.fromarray(img_array)
        result = result.filter(ImageFilter.GaussianBlur(radius=0.3))
        return result
    
    def create_paper_texture(self, width, height):
        """Tạo texture giấy"""
        paper = Image.new('RGB', (width, height), color=(252, 252, 248))
        texture_array = np.array(paper)
        paper_noise = np.random.normal(0, 3, texture_array.shape)
        texture_array = np.clip(texture_array + paper_noise, 0, 255).astype(np.uint8)
        return Image.fromarray(texture_array)
    
    def create_parchment_texture(self, width, height):
        """Tạo texture giấy da/cũ."""
        # Màu nền vàng của giấy cũ
        paper = Image.new('RGB', (width, height), color=(224, 212, 184))
        texture_array = np.array(paper).astype(np.float64)
        
        # Thêm nhiễu (noise) để tạo độ sần
        paper_noise = np.random.normal(0, 8, texture_array.shape)
        texture_array = np.clip(texture_array + paper_noise, 0, 255)
        
        paper_with_noise = Image.fromarray(texture_array.astype(np.uint8)).convert('RGBA')

        # Tạo một lớp trong suốt để vẽ các vết ố
        blotch_layer = Image.new('RGBA', (width, height), (0,0,0,0))
        blotch_draw = ImageDraw.Draw(blotch_layer)

        for _ in range(25): # Số lượng vết ố
            x1, y1 = random.randint(-50, width), random.randint(-50, height)
            x2, y2 = x1 + random.randint(100, 300), y1 + random.randint(100, 250)
            blotch_color = (80, 50, 30, random.randint(8, 20)) # Màu nâu, rất trong suốt
            blotch_draw.ellipse([(x1,y1), (x2,y2)], fill=blotch_color)
        
        # Làm mờ các vết ố để chúng trông mềm mại hơn
        blotch_layer = blotch_layer.filter(ImageFilter.GaussianBlur(radius=25))

        # Kết hợp giấy sần và lớp vết ố
        combined = Image.alpha_composite(paper_with_noise, blotch_layer)
        result = combined.filter(ImageFilter.GaussianBlur(radius=0.5)).convert('RGB')
        return result
    
    def create_olined_texture(self, width, height, margin, line_height):
        """Tạo giấy 4 ô ly Việt Nam (4 dòng nhỏ, 1 dòng đậm)"""
        paper = Image.new('RGB', (width, height), color=(255, 255, 255))
        draw = ImageDraw.Draw(paper)
        
        # Vẽ các dòng ngang
        for i in range(margin, height - margin, line_height // 4):
            # Dòng đậm (dòng chính) - màu xanh đậm
            if (i - margin) % line_height == 0:
                color = (100, 100, 200)
                width_line = 2
            # Dòng mỏng (dòng phụ) - màu xanh nhạt
            else:
                color = (180, 180, 255)
                width_line = 1
            
            draw.line([(margin, i), (width - margin, i)], fill=color, width=width_line)
        
        # Vẽ dòng kẻ dọc (ô ly)
        # Khoảng cách giữa các cột ô ly
        column_width = 80
        for x in range(margin + 60, width - margin, column_width):
            # Dòng dọc đậm hơn cho viền cột
            if (x - margin - 60) % (column_width * 2) == 0:
                draw.line([(x, margin), (x, height - margin)], fill=(150, 150, 200), width=2)
            else:
                draw.line([(x, margin), (x, height - margin)], fill=(200, 200, 220), width=1)
        
        # Vẽ lề đỏ bên trái (như giấy học sinh)
        margin_line_x = margin + 50
        draw.line([(margin_line_x, margin), (margin_line_x, height - margin)], 
                 fill=(255, 100, 100), width=2)
        
        return paper

    def create_exam_texture(self, width, height, margin, line_height):
        """Tạo giấy kiểm tra (dòng đỏ, lề xanh)"""
        paper = Image.new('RGB', (width, height), color=(255, 255, 255))
        draw = ImageDraw.Draw(paper)
        # Dòng đỏ
        for i in range(margin + line_height, height - margin, line_height):
            draw.line([(margin, i), (width - margin, i)], fill=(255, 100, 100), width=1)
        # Lề xanh
        draw.line([(margin + 50, margin), (margin + 50, height - margin)], fill=(100, 180, 255), width=2)
        return paper

    def create_calligraphy_texture(self, width, height, margin, line_height):
        """Tạo giấy thư pháp (nền vàng nhạt, lưới caro mờ)"""
        paper = Image.new('RGB', (width, height), color=(255, 250, 220))
        draw = ImageDraw.Draw(paper)
        # Lưới caro
        for i in range(margin, height - margin, line_height):
            draw.line([(margin, i), (width - margin, i)], fill=(220, 200, 150), width=1)
        for x in range(margin + 60, width - margin, 60):
            draw.line([(x, margin), (x, height - margin)], fill=(220, 200, 150), width=1)
        return paper

    def get_canvas_size(self, paper_size, custom_width, custom_height, num_lines, line_height, margin):
        if paper_size == "A4":
            # 794x1123 px ~ 210x297mm ở 96dpi
            return 794, max(1123, num_lines * line_height + margin * 2)
        elif paper_size == "A5":
            # 559x794 px ~ 148x210mm ở 96dpi
            return 559, max(794, num_lines * line_height + margin * 2)
        elif paper_size == "Custom":
            return custom_width, custom_height
        else:
            return 800, max(600, num_lines * line_height + margin * 2)

    def generate_handwriting(self, text, output_path=None, 
                           font_size=28, line_spacing=1.5, margin=50,
                           ink_color=(25, 25, 112), paper_style="lined", selected_font=None, progress_callback=None,
                           paper_size="A4", custom_width=800, custom_height=600, black_white=False):
        """Tạo chữ viết tay từ text"""
        import tkinter.messagebox as tkmsg
        if not self.fonts:
            raise Exception("Không có font nào khả dụng!")
        
        # Chọn font
        fallback_used = False
        if selected_font and selected_font in self.fonts:
            font_path = selected_font
        else:
            font_path = random.choice(self.fonts)
        
        try:
            if font_path:
                font = ImageFont.truetype(font_path, font_size, encoding='unic')
                # Kiểm tra font có vẽ được ký tự tiếng Việt không
                test_text = "ăâđêôơưàáảãạĐ"
                bbox = font.getbbox(test_text)
                if bbox[2] <= 0 or bbox[3] <= 0:
                    raise ValueError("Font không vẽ được bounding box cho ký tự TV.")
            else:
                font = ImageFont.load_default()
                fallback_used = True
        except Exception as e:
            print(f"Lỗi load font {font_path}: {e}")
            font = ImageFont.load_default()
            fallback_used = True
        
        if fallback_used:
            try:
                tkmsg.showwarning("Cảnh báo font", "Font bạn chọn bị lỗi hoặc không hỗ trợ tiếng Việt. Đã chuyển sang font mặc định.")
            except Exception:
                pass
        
        # Xử lý text với encoding UTF-8
        if isinstance(text, bytes):
            text = text.decode('utf-8')
        
        # Tính toán kích thước canvas trước
        line_height = int(font_size * line_spacing)
        canvas_width, canvas_height = self.get_canvas_size(paper_size, custom_width, custom_height, 1, line_height, margin)
        
        # Tính toán số ký tự tối đa trên mỗi dòng dựa trên kích thước giấy thực tế
        available_width = canvas_width - margin * 2 - 60  # Trừ lề và khoảng cách bắt đầu
        # Ước tính số ký tự dựa trên font size (trung bình 1 ký tự = font_size * 0.6)
        estimated_char_width = font_size * 0.6
        max_chars_per_line = int(available_width / estimated_char_width)
        
        # Sử dụng số ký tự đã tính toán thay vì cố định 70
        wrapper = textwrap.TextWrapper(width=max_chars_per_line)
        lines = []
        for paragraph in text.split('\n'):
            if paragraph.strip():
                lines.extend(wrapper.wrap(paragraph))
            else:
                lines.append('')
        
        # Tính lại chiều cao canvas dựa trên số dòng thực tế
        canvas_width, canvas_height = self.get_canvas_size(paper_size, custom_width, custom_height, len(lines), line_height, margin)
        
        # Tạo ảnh nền dựa trên kiểu giấy
        if paper_style == "parchment":
            image = self.create_parchment_texture(canvas_width, canvas_height)
        elif paper_style == "olined":
            image = self.create_olined_texture(canvas_width, canvas_height, margin, line_height)
        elif paper_style == "exam":
            image = self.create_exam_texture(canvas_width, canvas_height, margin, line_height)
        elif paper_style == "calligraphy":
            image = self.create_calligraphy_texture(canvas_width, canvas_height, margin, line_height)
        elif paper_style == "grid":
            image = self.create_paper_texture(canvas_width, canvas_height)
            draw = ImageDraw.Draw(image)
            # Vẽ lưới caro
            for i in range(margin, canvas_height - margin, line_height):
                draw.line([(margin, i), (canvas_width - margin, i)], fill=(220, 220, 220), width=1)
            for x in range(margin + 60, canvas_width - margin, 60):
                draw.line([(x, margin), (x, canvas_height - margin)], fill=(220, 220, 220), width=1)
        elif paper_style == "dotted":
            image = self.create_paper_texture(canvas_width, canvas_height)
            draw = ImageDraw.Draw(image)
            # Vẽ chấm bi
            for i in range(margin, canvas_height - margin, line_height):
                for x in range(margin + 60, canvas_width - margin, 60):
                    draw.ellipse((x-1, i-1, x+1, i+1), fill=(180, 180, 180))
        elif paper_style == "blank":
            image = self.create_paper_texture(canvas_width, canvas_height)
        else:
            image = self.create_paper_texture(canvas_width, canvas_height)
        
        draw = ImageDraw.Draw(image)
        
        if paper_style == "lined":
            line_color = (200, 200, 255, 50)
            for i in range(margin + line_height, canvas_height - margin, line_height):
                draw.line([(margin, i), (canvas_width - margin, i)], 
                         fill=line_color, width=1)
        
        if paper_style == "lined":
            margin_line_x = margin + 50
            draw.line([(margin_line_x, margin), (margin_line_x, canvas_height - margin)], 
                     fill=(255, 200, 200, 100), width=1)
        
        y_position = margin + line_height
        
        for idx, line in enumerate(lines):
            if line.strip():
                # Tăng độ ngẫu nhiên cho vị trí bắt đầu dòng
                x_position = margin + 60 + random.randint(-8, 20)
                
                # Tăng độ nghiêng dòng (slant) để tự nhiên hơn
                line_slant = random.randint(-3, 3)
                y_with_slant = y_position + line_slant
                
                # Thêm hiệu ứng lệch dòng ngẫu nhiên
                line_wobble = random.randint(-2, 2)
                y_with_slant += line_wobble
                
                words = line.split()
                current_x = x_position
                
                for word_idx, word in enumerate(words):
                    if word_idx > 0:
                        # Tăng khoảng cách giữa các từ
                        current_x += random.randint(10, 20)
                    
                    # Kiểm tra xem từ có vượt quá biên phải không
                    bbox = draw.textbbox((0, 0), word, font=font)
                    word_width = bbox[2] - bbox[0]
                    if current_x + word_width > canvas_width - margin:
                        # Xuống dòng mới với vị trí ngẫu nhiên
                        y_position += line_height
                        current_x = margin + 60 + random.randint(-8, 20)
                        line_slant = random.randint(-3, 3)
                        y_with_slant = y_position + line_slant + random.randint(-2, 2)
                    
                    self.add_natural_variations(draw, word, current_x, 
                                              y_with_slant, font, ink_color)
                    
                    current_x += word_width
            
            y_position += line_height
            if progress_callback is not None:
                progress_callback(idx + 1)
        
        image = self.add_ink_effects(image)
        
        # Chuyển đổi sang đen trắng nếu được yêu cầu
        if black_white:
            image = image.convert('L')  # Chuyển sang grayscale
            # Tăng độ tương phản để giống scan
            from PIL import ImageEnhance
            enhancer = ImageEnhance.Contrast(image)
            image = enhancer.enhance(1.5)  # Tăng độ tương phản 50%
            # Chuyển về RGB để lưu PNG
            image = image.convert('RGB')
        
        if output_path:
            image.save(output_path, 'PNG', quality=95)
        
        return image

class HandwritingGUI:
    def __init__(self, root):
        self.font_lang = 'english'
        self.generator = HandwritingGenerator(font_lang=self.font_lang)
        self.lang = 'vi'
        self.L = LANGS[self.lang]
        self.current_image = None
        self.zoom_level = 100  # Mức zoom ban đầu là 100%
        self.settings = self.load_settings()
        self.progress = None  # Progress bar
        self.root = root
        self.root.title(self.L['title'])
        self.root.geometry("1200x800")
        self.root.configure(bg='#f0f0f0')
        # Bỏ combobox chọn ngôn ngữ font, chỉ giữ label hướng dẫn
        self.font_lang_label = ttk.Label(self.root, text="Chọn font (có ghi chú nước):")
        self.font_lang_label.place(x=10, y=5)
        self.font_count_label = ttk.Label(self.root, text="")
        self.font_count_label.place(x=340, y=5)
        self.font_hint_label = ttk.Label(self.root, text="", foreground="red")
        self.font_hint_label.place(x=150, y=30)
        self.font_file_label = ttk.Label(self.root, text="", foreground="blue")
        self.font_file_label.place(x=150, y=55)
        self.setup_ui()
        self.load_settings_to_ui()
        self.setup_font_list()
        self.refresh_template_list()  # Load danh sách template

    def setup_ui(self):
        """Thiết lập giao diện người dùng"""
        
        # Main frame
        main_frame = ttk.Frame(self.root, padding="10")
        main_frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        
        # Configure grid weights
        self.root.columnconfigure(0, weight=1)
        self.root.rowconfigure(0, weight=1)
        main_frame.columnconfigure(1, weight=1)
        main_frame.rowconfigure(1, weight=1)
        
        # Title
        title_label = ttk.Label(main_frame, text=self.L['title'], 
                               font=('Arial', 16, 'bold'))
        title_label.grid(row=0, column=0, columnspan=3, pady=(0, 10))
        
        # Left panel - Settings (với scrollbar)
        settings_container = ttk.Frame(main_frame)
        settings_container.grid(row=1, column=0, sticky=(tk.W, tk.E, tk.N, tk.S), padx=(0, 10))
        settings_container.columnconfigure(0, weight=1)
        settings_container.rowconfigure(0, weight=1)
        
        # Tạo canvas và scrollbar cho settings
        settings_canvas = tk.Canvas(settings_container, width=350)
        settings_scrollbar = ttk.Scrollbar(settings_container, orient="vertical", command=settings_canvas.yview)
        settings_frame = ttk.LabelFrame(settings_canvas, text="⚙️ Cài đặt", padding="10")
        
        # Configure canvas
        settings_canvas.configure(yscrollcommand=settings_scrollbar.set)
        settings_canvas.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        settings_scrollbar.grid(row=0, column=1, sticky=(tk.N, tk.S))
        
        # Tạo window trong canvas cho settings_frame
        settings_canvas.create_window((0, 0), window=settings_frame, anchor="nw")
        
        # Configure settings_frame để mở rộng theo nội dung
        settings_frame.bind("<Configure>", lambda e: settings_canvas.configure(scrollregion=settings_canvas.bbox("all")))
        
        # Bind mouse wheel cho settings_canvas
        def _on_mousewheel(event):
            settings_canvas.yview_scroll(int(-1*(event.delta/120)), "units")
        settings_canvas.bind_all("<MouseWheel>", _on_mousewheel)
        
        # Font size
        ttk.Label(settings_frame, text=self.L['font_size']).grid(row=0, column=0, sticky=tk.W, pady=2)
        self.font_size_var = tk.IntVar(value=28)
        font_size_scale = ttk.Scale(settings_frame, from_=16, to=48, 
                                   variable=self.font_size_var, orient=tk.HORIZONTAL)
        font_size_scale.grid(row=0, column=1, sticky=(tk.W, tk.E), pady=2)
        self.font_size_label = ttk.Label(settings_frame, text="28")
        self.font_size_label.grid(row=0, column=2, pady=2)
        font_size_scale.configure(command=self.update_font_size_label)
        
        # Font selection
        ttk.Label(settings_frame, text=self.L['choose_font']).grid(row=1, column=0, sticky=tk.W, pady=2)
        self.selected_font = tk.StringVar()
        self.font_combo = ttk.Combobox(settings_frame, textvariable=self.selected_font, 
                                      state="readonly", width=25)
        self.font_combo.grid(row=1, column=1, columnspan=2, sticky=(tk.W, tk.E), pady=2)
        self.font_combo.bind('<<ComboboxSelected>>', self.update_font_demo)
        
        # Font demo
        ttk.Label(settings_frame, text=self.L['font_demo']).grid(row=2, column=0, sticky=tk.W, pady=2)
        demo_frame = ttk.Frame(settings_frame, height=60, width=300)
        demo_frame.grid(row=2, column=1, columnspan=2, sticky=(tk.W, tk.E), pady=2)
        demo_frame.grid_propagate(False)
        self.font_demo_label = ttk.Label(demo_frame, text="Xin chào! Hello!", font=('Arial', 10), foreground='blue')
        self.font_demo_label.place(relx=0.5, rely=0.5, anchor='center')
        # Thêm ô nhập text demo
        self.demo_text_var = tk.StringVar(value="Xin chào! Hello! 123")
        self.demo_text_entry = ttk.Entry(settings_frame, textvariable=self.demo_text_var, width=30)
        self.demo_text_entry.grid(row=3, column=1, columnspan=2, sticky=(tk.W, tk.E), pady=(0, 5))
        self.demo_text_entry.bind('<KeyRelease>', self.update_font_demo)
        
        # Line spacing
        ttk.Label(settings_frame, text=self.L['line_spacing']).grid(row=4, column=0, sticky=tk.W, pady=2)
        self.line_spacing_var = tk.DoubleVar(value=1.5)
        line_spacing_scale = ttk.Scale(settings_frame, from_=1.0, to=3.0, 
                                      variable=self.line_spacing_var, orient=tk.HORIZONTAL)
        line_spacing_scale.grid(row=4, column=1, sticky=(tk.W, tk.E), pady=2)
        self.line_spacing_label = ttk.Label(settings_frame, text="1.5")
        self.line_spacing_label.grid(row=4, column=2, pady=2)
        line_spacing_scale.configure(command=self.update_line_spacing_label)
        
        # Margin
        ttk.Label(settings_frame, text=self.L['margin']).grid(row=5, column=0, sticky=tk.W, pady=2)
        self.margin_var = tk.IntVar(value=50)
        margin_scale = ttk.Scale(settings_frame, from_=20, to=100, 
                                variable=self.margin_var, orient=tk.HORIZONTAL)
        margin_scale.grid(row=5, column=1, sticky=(tk.W, tk.E), pady=2)
        self.margin_label = ttk.Label(settings_frame, text="50")
        self.margin_label.grid(row=5, column=2, pady=2)
        margin_scale.configure(command=self.update_margin_label)
        
        # Ink color
        ttk.Label(settings_frame, text=self.L['ink_color']).grid(row=6, column=0, sticky=tk.W, pady=2)
        self.ink_color = (25, 25, 112)  # Default blue
        self.color_button = tk.Button(settings_frame, text="Chọn màu", 
                                     bg=self.rgb_to_hex(self.ink_color), 
                                     command=self.choose_color)
        self.color_button.grid(row=6, column=1, sticky=(tk.W, tk.E), pady=2)
        
        # Paper style
        ttk.Label(settings_frame, text=self.L['paper_style']).grid(row=7, column=0, sticky=tk.W, pady=2)
        self.paper_style_var = tk.StringVar(value="lined")
        paper_combo = ttk.Combobox(settings_frame, textvariable=self.paper_style_var,
                                  values=["lined", "blank", "grid", "dotted", "parchment", "olined", "exam", "calligraphy"], 
                                  state="readonly")
        paper_combo.grid(row=7, column=1, sticky=(tk.W, tk.E), pady=2)
        
        # Thêm tooltip cho các kiểu giấy
        paper_style_tooltips = {
            "lined": "Giấy kẻ ngang",
            "blank": "Giấy trắng",
            "grid": "Giấy ô vuông",
            "dotted": "Giấy chấm bi",
            "parchment": "Giấy giả cổ",
            "olined": "Giấy 4 ô ly Việt Nam",
            "exam": "Giấy kiểm tra",
            "calligraphy": "Giấy thư pháp"
        }
        
        def show_paper_tooltip(event):
            selected = self.paper_style_var.get()
            if selected in paper_style_tooltips:
                tooltip_text = paper_style_tooltips[selected]
                # Hiển thị tooltip trong status bar
                self.status_bar.config(text=f"Kiểu giấy: {tooltip_text}")
        
        paper_combo.bind('<<ComboboxSelected>>', show_paper_tooltip)
        
        # Black & White option
        self.black_white_var = tk.BooleanVar(value=False)
        black_white_check = ttk.Checkbutton(settings_frame, text="Xuất đen trắng (như scan)", 
                                           variable=self.black_white_var)
        black_white_check.grid(row=7, column=2, sticky=tk.W, pady=2)
        
        # Template System
        template_frame = ttk.LabelFrame(settings_frame, text="📋 Template", padding="5")
        template_frame.grid(row=8, column=0, columnspan=3, sticky=(tk.W, tk.E), pady=5)
        
        # Template name input
        ttk.Label(template_frame, text="Tên template:").grid(row=0, column=0, sticky=tk.W, pady=2)
        self.template_name_var = tk.StringVar()
        template_name_entry = ttk.Entry(template_frame, textvariable=self.template_name_var, width=20)
        template_name_entry.grid(row=0, column=1, sticky=(tk.W, tk.E), pady=2)
        
        # Template buttons
        template_buttons_frame = ttk.Frame(template_frame)
        template_buttons_frame.grid(row=1, column=0, columnspan=2, sticky=(tk.W, tk.E), pady=2)
        
        ttk.Button(template_buttons_frame, text="💾 Lưu template", 
                  command=self.save_template).pack(side=tk.LEFT, padx=2)
        ttk.Button(template_buttons_frame, text="📂 Tải template", 
                  command=self.load_template).pack(side=tk.LEFT, padx=2)
        ttk.Button(template_buttons_frame, text="🗑️ Xóa template", 
                  command=self.delete_template).pack(side=tk.LEFT, padx=2)
        
        # Template list
        ttk.Label(template_frame, text="Template đã lưu:").grid(row=2, column=0, sticky=tk.W, pady=(5,2))
        self.template_listbox = tk.Listbox(template_frame, height=3)
        self.template_listbox.grid(row=3, column=0, columnspan=2, sticky=(tk.W, tk.E), pady=2)
        self.template_listbox.bind('<Double-Button-1>', self.load_template_from_list)
        
        # Configure template frame column weights
        template_frame.columnconfigure(1, weight=1)
        
        # Paper size
        ttk.Label(settings_frame, text=self.L['paper_size']).grid(row=9, column=0, sticky=tk.W, pady=2)
        self.paper_size_var = tk.StringVar(value="A4")
        paper_size_combo = ttk.Combobox(settings_frame, textvariable=self.paper_size_var,
            values=["A4", "A5", "Custom"], state="readonly", width=10)
        paper_size_combo.grid(row=9, column=1, sticky=(tk.W, tk.E), pady=2)
        self.custom_width_var = tk.IntVar(value=800)
        self.custom_height_var = tk.IntVar(value=600)
        self.custom_size_frame = ttk.Frame(settings_frame)
        self.custom_size_frame.grid(row=9, column=2, sticky=(tk.W, tk.E), pady=2)
        ttk.Label(self.custom_size_frame, text="W:").pack(side=tk.LEFT)
        self.custom_width_entry = ttk.Entry(self.custom_size_frame, textvariable=self.custom_width_var, width=5)
        self.custom_width_entry.pack(side=tk.LEFT)
        ttk.Label(self.custom_size_frame, text="H:").pack(side=tk.LEFT)
        self.custom_height_entry = ttk.Entry(self.custom_size_frame, textvariable=self.custom_height_var, width=5)
        self.custom_height_entry.pack(side=tk.LEFT)
        self.custom_size_frame.grid_remove()
        paper_size_combo.bind('<<ComboboxSelected>>', self.on_paper_size_change)
        
        # Configure column weights for settings frame
        settings_frame.columnconfigure(1, weight=1)
        
        # Dark Mode Toggle
        dark_mode_frame = ttk.Frame(settings_frame)
        dark_mode_frame.grid(row=10, column=0, columnspan=3, sticky=(tk.W, tk.E), pady=5)
        
        self.dark_mode_var = tk.BooleanVar(value=False)
        dark_mode_check = ttk.Checkbutton(dark_mode_frame, text="🌙 Dark Mode", 
                                         variable=self.dark_mode_var, command=self.toggle_dark_mode)
        dark_mode_check.pack(side=tk.LEFT, padx=5)
        
        # Action buttons
        button_frame = ttk.Frame(settings_frame)
        button_frame.grid(row=11, column=0, columnspan=3, pady=10, sticky=(tk.W, tk.E))
        
        ttk.Button(button_frame, text=self.L['create'], 
                  command=self.generate_handwriting).pack(fill=tk.X, pady=2)
        ttk.Button(button_frame, text=self.L['save_settings'], 
                  command=self.save_settings).pack(fill=tk.X, pady=2)
        ttk.Button(button_frame, text=self.L['open_file'], 
                  command=self.load_file).pack(fill=tk.X, pady=2)
        
        # Export buttons
        export_frame = ttk.LabelFrame(settings_frame, text="📤 Xuất file", padding="5")
        export_frame.grid(row=12, column=0, columnspan=3, pady=10, sticky=(tk.W, tk.E))
        
        ttk.Button(export_frame, text=self.L['export_png'], 
                  command=self.export_png).pack(fill=tk.X, pady=1)
        ttk.Button(export_frame, text=self.L['export_pdf'], 
                  command=self.export_pdf).pack(fill=tk.X, pady=1)
        ttk.Button(export_frame, text=self.L['export_word'], 
                  command=self.export_word).pack(fill=tk.X, pady=1)
        ttk.Button(export_frame, text="📐 SVG", 
                  command=self.export_svg).pack(fill=tk.X, pady=1)
        
        # Middle panel - Text input (cải tiến với scrollbar ngang)
        text_frame = ttk.LabelFrame(main_frame, text=self.L['input_text'], padding="10")
        text_frame.grid(row=1, column=1, sticky=(tk.W, tk.E, tk.N, tk.S), padx=(0, 10))
        text_frame.columnconfigure(0, weight=1)
        text_frame.rowconfigure(0, weight=1)
        
        # Text area container với scrollbar ngang
        text_container = ttk.Frame(text_frame)
        text_container.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        text_container.columnconfigure(0, weight=1)
        text_container.rowconfigure(0, weight=1)
        
        # Text area với cả scrollbar dọc và ngang
        self.text_area = tk.Text(text_container, wrap=tk.NONE, font=('Arial', 11), height=15)
        text_scrollbar_v = ttk.Scrollbar(text_container, orient="vertical", command=self.text_area.yview)
        text_scrollbar_h = ttk.Scrollbar(text_container, orient="horizontal", command=self.text_area.xview)
        
        # Configure text area
        self.text_area.configure(yscrollcommand=text_scrollbar_v.set, xscrollcommand=text_scrollbar_h.set)
        
        # Grid layout
        self.text_area.grid(row=0, column=0, sticky="nsew")
        text_scrollbar_v.grid(row=0, column=1, sticky="ns")
        text_scrollbar_h.grid(row=1, column=0, sticky="ew")
        
        # Text wrap toggle
        wrap_frame = ttk.Frame(text_frame)
        wrap_frame.grid(row=1, column=0, sticky=(tk.W, tk.E), pady=(5, 0))
        
        self.wrap_var = tk.BooleanVar(value=False)
        wrap_check = ttk.Checkbutton(wrap_frame, text="Tự động xuống dòng", 
                                   variable=self.wrap_var, command=self.toggle_text_wrap)
        wrap_check.pack(side=tk.LEFT)
        
        # Character count
        self.char_count_label = ttk.Label(text_frame, text=f"{self.L['char_count']} 0")
        self.char_count_label.grid(row=2, column=0, sticky=tk.W, pady=(5, 0))
        
        # Sample text
        sample_text = """Xin chào! Đây là công cụ chuyển đổi văn bản thành chữ viết tay.

Bạn có thể:
• Nhập văn bản bất kỳ vào ô này
• Tùy chỉnh kích thước chữ, màu mực
• Chọn kiểu giấy (kẻ ngang, trắng, ô vuông)
• Xuất ra file PNG, PDF, hoặc Word

Tool hỗ trợ tiếng Việt đầy đủ với các dấu: à, á, ả, ã, ạ, ă, ằ, ắ, ẳ, ẵ, ặ, â, ầ, ấ, ẩ, ẫ, ậ, đ, è, é, ẻ, ẽ, ẹ, ê, ề, ế, ể, ễ, ệ, ì, í, ỉ, ĩ, ị, ò, ó, ỏ, õ, ọ, ô, ồ, ố, ổ, ỗ, ộ, ơ, ờ, ớ, ở, ỡ, ợ, ù, ú, ủ, ũ, ụ, ư, ừ, ứ, ử, ữ, ự, ỳ, ý, ỷ, ỹ, ỵ.

Chúc bạn sử dụng vui vẻ! 😊

Các số: 0123456789
Các ký tự đặc biệt: !@#$%^&*()_+-=[]{}|;':",./<>?`~"""
        
        self.text_area.insert('1.0', sample_text)
        self.text_area.bind('<KeyRelease>', self.update_char_count)
        self.update_char_count()
        
        # Right panel - Preview
        preview_frame = ttk.LabelFrame(main_frame, text=self.L['preview'], padding="10")
        preview_frame.grid(row=1, column=2, sticky=(tk.W, tk.E, tk.N, tk.S))
        preview_frame.columnconfigure(0, weight=1)
        preview_frame.rowconfigure(1, weight=1) # Thay đổi để chừa chỗ cho thanh zoom
        
        # --- Thanh công cụ Zoom ---
        zoom_controls_frame = ttk.Frame(preview_frame)
        zoom_controls_frame.grid(row=0, column=0, sticky=(tk.W, tk.E), pady=(0, 5))

        zoom_out_button = ttk.Button(zoom_controls_frame, text="➖", command=self.zoom_out, width=4)
        zoom_out_button.pack(side=tk.LEFT, padx=2)

        self.zoom_label = ttk.Label(zoom_controls_frame, text=f"{self.zoom_level}%", width=6, anchor="center")
        self.zoom_label.pack(side=tk.LEFT, padx=5)

        zoom_in_button = ttk.Button(zoom_controls_frame, text="➕", command=self.zoom_in, width=4)
        zoom_in_button.pack(side=tk.LEFT, padx=2)
        
        # Preview canvas with scrollbar (cải tiến)
        canvas_frame = ttk.Frame(preview_frame)
        canvas_frame.grid(row=1, column=0, sticky=(tk.W, tk.E, tk.N, tk.S)) # Chuyển xuống dòng 1
        canvas_frame.columnconfigure(0, weight=1)
        canvas_frame.rowconfigure(0, weight=1)
        
        self.preview_canvas = tk.Canvas(canvas_frame, bg='white', width=400, height=500)
        
        # Thêm cả hai thanh cuộn: dọc và ngang
        preview_scrollbar_v = ttk.Scrollbar(canvas_frame, orient="vertical", command=self.preview_canvas.yview)
        preview_scrollbar_h = ttk.Scrollbar(canvas_frame, orient="horizontal", command=self.preview_canvas.xview)
        self.preview_canvas.configure(yscrollcommand=preview_scrollbar_v.set, xscrollcommand=preview_scrollbar_h.set)
        
        self.preview_canvas.grid(row=0, column=0, sticky="nsew")
        preview_scrollbar_v.grid(row=0, column=1, sticky="ns")
        preview_scrollbar_h.grid(row=1, column=0, sticky="ew")

        # --- Gán sự kiện để di chuyển ảnh (panning) bằng chuột ---
        self.preview_canvas.bind("<ButtonPress-1>", self.on_canvas_press)
        self.preview_canvas.bind("<B1-Motion>", self.on_canvas_drag)
        self.preview_canvas.bind("<Enter>", self.on_canvas_enter)
        self.preview_canvas.bind("<Leave>", self.on_canvas_leave)
        
        # Thêm keyboard shortcuts cho preview
        self.preview_canvas.bind("<KeyPress-Left>", lambda e: self.preview_canvas.xview_scroll(-1, "units"))
        self.preview_canvas.bind("<KeyPress-Right>", lambda e: self.preview_canvas.xview_scroll(1, "units"))
        self.preview_canvas.bind("<KeyPress-Up>", lambda e: self.preview_canvas.yview_scroll(-1, "units"))
        self.preview_canvas.bind("<KeyPress-Down>", lambda e: self.preview_canvas.yview_scroll(1, "units"))
        self.preview_canvas.bind("<KeyPress-Home>", lambda e: self.preview_canvas.xview_moveto(0))
        self.preview_canvas.bind("<KeyPress-End>", lambda e: self.preview_canvas.xview_moveto(1))
        self.preview_canvas.focus_set()  # Cho phép nhận keyboard events
        
        # Status bar
        self.status_bar = ttk.Label(main_frame, text=self.L['ready'], relief=tk.SUNKEN)
        self.status_bar.grid(row=2, column=0, columnspan=3, sticky=(tk.W, tk.E), pady=(10, 0))
        # Progress bar (ẩn mặc định)
        self.progress = ttk.Progressbar(main_frame, orient="horizontal", mode="determinate", length=300)
        self.progress.grid(row=3, column=0, columnspan=3, pady=(2, 0))
        self.progress.grid_remove()
        
    def update_font_size_label(self, value):
        """Cập nhật label kích thước font"""
        self.font_size_label.config(text=str(int(float(value))))
        
    def update_line_spacing_label(self, value):
        """Cập nhật label khoảng cách dòng"""
        self.line_spacing_label.config(text=f"{float(value):.1f}")
        
    def update_margin_label(self, value):
        """Cập nhật label lề"""
        self.margin_label.config(text=str(int(float(value))))
        
    def update_char_count(self, event=None):
        """Cập nhật số ký tự"""
        text = self.text_area.get('1.0', tk.END)
        char_count = len(text) - 1  # Trừ ký tự newline cuối
        self.char_count_label.config(text=f"{self.L['char_count']} {char_count}")
    
    def toggle_text_wrap(self):
        """Chuyển đổi chế độ wrap cho text area"""
        if self.wrap_var.get():
            self.text_area.configure(wrap=tk.WORD)
        else:
            self.text_area.configure(wrap=tk.NONE)
    
    def rgb_to_hex(self, rgb):
        """Chuyển RGB thành hex"""
        return f"#{rgb[0]:02x}{rgb[1]:02x}{rgb[2]:02x}"
        
    def choose_color(self):
        """Chọn màu mực"""
        color = colorchooser.askcolor(initialcolor=self.rgb_to_hex(self.ink_color))
        if color[0]:  # Nếu user chọn màu
            self.ink_color = tuple(int(c) for c in color[0])
            self.color_button.config(bg=color[1])
            
    def load_file(self):
        """Mở file text"""
        file_path = filedialog.askopenfilename(
            title="Chọn file text",
            filetypes=[("Text files", "*.txt"), ("All files", "*.*")]
        )
        if file_path:
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
                self.text_area.delete('1.0', tk.END)
                self.text_area.insert('1.0', content)
                self.update_char_count()
                self.status_bar.config(text=f"Đã mở: {os.path.basename(file_path)}")
            except Exception as e:
                messagebox.showerror("Lỗi", f"Không thể mở file: {e}")
                
    def generate_handwriting(self):
        """Tạo chữ viết tay"""
        text = self.text_area.get('1.0', tk.END).strip()
        if not text:
            messagebox.showwarning("Cảnh báo", "Vui lòng nhập văn bản!")
            return
        self.status_bar.config(text=self.L['generating'])
        self.root.update()
        self.zoom_level = 100 # Reset zoom về 100% mỗi khi tạo ảnh mới
        self.zoom_label.config(text=f"{self.zoom_level}%")
        def generate_thread():
            try:
                selected_display_name = self.selected_font.get()
                font_path, font_lang = None, None
                if hasattr(self, 'font_lang_map') and selected_display_name in self.font_lang_map:
                    font_path, font_lang = self.font_lang_map[selected_display_name]
                else:
                    font_path = None
                    font_lang = self.font_lang
                wrapper = textwrap.TextWrapper(width=70)
                lines = []
                for paragraph in text.split('\n'):
                    if paragraph.strip():
                        lines.extend(wrapper.wrap(paragraph))
                    else:
                        lines.append('')
                total_lines = len(lines)
                self.root.after(0, lambda: self.show_progress(total_lines))
                paper_size = self.paper_size_var.get()
                custom_width = self.custom_width_var.get()
                custom_height = self.custom_height_var.get()
                # Tạo generator đúng ngôn ngữ font
                temp_gen = HandwritingGenerator(font_lang=font_lang)
                self.current_image = temp_gen.generate_handwriting(
                    text=text,
                    font_size=self.font_size_var.get(),
                    line_spacing=self.line_spacing_var.get(),
                    margin=self.margin_var.get(),
                    ink_color=self.ink_color,
                    paper_style=self.paper_style_var.get(),
                    selected_font=font_path,
                    progress_callback=self.update_progress,
                    paper_size=paper_size,
                    custom_width=custom_width,
                    custom_height=custom_height,
                    black_white=self.black_white_var.get()
                )
                self.root.after(0, self.update_preview)
                self.root.after(0, lambda: self.status_bar.config(text="Tạo thành công!"))
            except Exception as e:
                self.root.after(0, lambda: messagebox.showerror("Lỗi", f"Không thể tạo: {e}"))
                self.root.after(0, lambda: self.status_bar.config(text="Lỗi tạo chữ viết tay"))
            finally:
                self.root.after(0, self.hide_progress)
        threading.Thread(target=generate_thread, daemon=True).start()

    def show_progress(self, maximum):
        self.progress['value'] = 0
        self.progress['maximum'] = maximum
        self.progress.grid()
        self.progress.update()
    def update_progress(self, value):
        self.progress['value'] = value
        self.progress.update()
    def hide_progress(self):
        self.progress.grid_remove()

    def update_preview(self):
        """Cập nhật preview với mức zoom hiện tại."""
        if self.current_image:
            # Tính toán kích thước mới dựa trên mức zoom
            zoom_factor = self.zoom_level / 100.0
            new_width = int(self.current_image.width * zoom_factor)
            new_height = int(self.current_image.height * zoom_factor)
            
            # Đảm bảo kích thước tối thiểu là 1x1 để tránh lỗi
            if new_width < 1: new_width = 1
            if new_height < 1: new_height = 1

            try:
                img_copy = self.current_image.copy()
                resized_img = img_copy.resize((new_width, new_height), Image.Resampling.LANCZOS)
                
                # Chuyển thành PhotoImage
                self.preview_photo = ImageTk.PhotoImage(resized_img)
                
                # Clear canvas và hiển thị image ở góc trên bên trái (NW)
                self.preview_canvas.delete("all")
                self.preview_canvas.create_image(0, 0, anchor=tk.NW, image=self.preview_photo)
                
                # Cập nhật scroll region để scrollbar hoạt động đúng
                self.preview_canvas.configure(scrollregion=self.preview_canvas.bbox("all"))
            except Exception as e:
                print(f"Lỗi khi cập nhật preview: {e}")

    def zoom_in(self):
        """Phóng to ảnh xem trước."""
        if self.zoom_level < 300: # Giới hạn mức zoom tối đa
            self.zoom_level += 20
            self.zoom_label.config(text=f"{self.zoom_level}%")
            self.update_preview()

    def zoom_out(self):
        """Thu nhỏ ảnh xem trước."""
        if self.zoom_level > 30: # Giới hạn mức zoom tối thiểu
            self.zoom_level -= 20
            self.zoom_label.config(text=f"{self.zoom_level}%")
            self.update_preview()
            
    def export_png(self):
        """Xuất file PNG"""
        if not self.current_image:
            messagebox.showwarning("Cảnh báo", "Vui lòng tạo chữ viết tay trước!")
            return
            
        file_path = filedialog.asksaveasfilename(
            title="Lưu file PNG",
            defaultextension=".png",
            filetypes=[("PNG files", "*.png"), ("All files", "*.*")]
        )
        if file_path:
            try:
                # Tạo lại ảnh với tùy chọn đen trắng nếu được chọn
                if self.black_white_var.get():
                    # Tạo lại ảnh với tùy chọn đen trắng
                    text = self.text_area.get('1.0', tk.END).strip()
                    selected_display_name = self.selected_font.get()
                    font_path, font_lang = None, None
                    if hasattr(self, 'font_lang_map') and selected_display_name in self.font_lang_map:
                        font_path, font_lang = self.font_lang_map[selected_display_name]
                    else:
                        font_path = None
                        font_lang = self.font_lang
                    
                    temp_gen = HandwritingGenerator(font_lang=font_lang)
                    export_image = temp_gen.generate_handwriting(
                        text=text,
                        font_size=self.font_size_var.get(),
                        line_spacing=self.line_spacing_var.get(),
                        margin=self.margin_var.get(),
                        ink_color=self.ink_color,
                        paper_style=self.paper_style_var.get(),
                        selected_font=font_path,
                        paper_size=self.paper_size_var.get(),
                        custom_width=self.custom_width_var.get(),
                        custom_height=self.custom_height_var.get(),
                        black_white=True
                    )
                    export_image.save(file_path, 'PNG', quality=95)
                else:
                    self.current_image.save(file_path, 'PNG', quality=95)
                
                messagebox.showinfo("Thành công", f"Đã lưu: {os.path.basename(file_path)}")
                self.status_bar.config(text=f"Đã xuất PNG: {os.path.basename(file_path)}")
            except Exception as e:
                messagebox.showerror("Lỗi", f"Không thể lưu file: {e}")
                
    def export_pdf(self):
        """Xuất file PDF"""
        if not self.current_image:
            messagebox.showwarning("Cảnh báo", "Vui lòng tạo chữ viết tay trước!")
            return
            
        file_path = filedialog.asksaveasfilename(
            title="Lưu file PDF",
            defaultextension=".pdf",
            filetypes=[("PDF files", "*.pdf"), ("All files", "*.*")]
        )
        if file_path:
            try:
                # Tạo lại ảnh với tùy chọn đen trắng nếu được chọn
                if self.black_white_var.get():
                    # Tạo lại ảnh với tùy chọn đen trắng
                    text = self.text_area.get('1.0', tk.END).strip()
                    selected_display_name = self.selected_font.get()
                    font_path, font_lang = None, None
                    if hasattr(self, 'font_lang_map') and selected_display_name in self.font_lang_map:
                        font_path, font_lang = self.font_lang_map[selected_display_name]
                    else:
                        font_path = None
                        font_lang = self.font_lang
                    
                    temp_gen = HandwritingGenerator(font_lang=font_lang)
                    export_image = temp_gen.generate_handwriting(
                        text=text,
                        font_size=self.font_size_var.get(),
                        line_spacing=self.line_spacing_var.get(),
                        margin=self.margin_var.get(),
                        ink_color=self.ink_color,
                        paper_style=self.paper_style_var.get(),
                        selected_font=font_path,
                        paper_size=self.paper_size_var.get(),
                        custom_width=self.custom_width_var.get(),
                        custom_height=self.custom_height_var.get(),
                        black_white=True
                    )
                else:
                    export_image = self.current_image
                
                # Chuyển PIL Image thành PDF
                img_buffer = io.BytesIO()
                export_image.save(img_buffer, format='PNG')
                img_buffer.seek(0)
                
                # Tạo PDF
                c = canvas.Canvas(file_path, pagesize=A4)
                
                # Tính toán kích thước để fit A4
                img_reader = ImageReader(img_buffer)
                img_width, img_height = export_image.size
                
                # Scale để fit A4 với margin
                page_width, page_height = A4
                margin = 50
                max_width = page_width - 2 * margin
                max_height = page_height - 2 * margin
                
                scale = min(max_width / img_width, max_height / img_height)
                new_width = img_width * scale
                new_height = img_height * scale
                
                # Vẽ image vào PDF
                x = (page_width - new_width) / 2
                y = (page_height - new_height) / 2
                
                c.drawImage(img_reader, x, y, new_width, new_height)
                c.save()
                
                messagebox.showinfo("Thành công", f"Đã lưu: {os.path.basename(file_path)}")
                self.status_bar.config(text=f"Đã xuất PDF: {os.path.basename(file_path)}")
            except Exception as e:
                messagebox.showerror("Lỗi", f"Không thể tạo PDF: {e}")
                
    def export_word(self):
        """Xuất file Word"""
        if not self.current_image:
            messagebox.showwarning("Cảnh báo", "Vui lòng tạo chữ viết tay trước!")
            return
            
        file_path = filedialog.asksaveasfilename(
            title="Lưu file Word",
            defaultextension=".docx",
            filetypes=[("Word files", "*.docx"), ("All files", "*.*")]
        )
        if file_path:
            try:
                # Tạo lại ảnh với tùy chọn đen trắng nếu được chọn
                if self.black_white_var.get():
                    # Tạo lại ảnh với tùy chọn đen trắng
                    text = self.text_area.get('1.0', tk.END).strip()
                    selected_display_name = self.selected_font.get()
                    font_path, font_lang = None, None
                    if hasattr(self, 'font_lang_map') and selected_display_name in self.font_lang_map:
                        font_path, font_lang = self.font_lang_map[selected_display_name]
                    else:
                        font_path = None
                        font_lang = self.font_lang
                    
                    temp_gen = HandwritingGenerator(font_lang=font_lang)
                    export_image = temp_gen.generate_handwriting(
                        text=text,
                        font_size=self.font_size_var.get(),
                        line_spacing=self.line_spacing_var.get(),
                        margin=self.margin_var.get(),
                        ink_color=self.ink_color,
                        paper_style=self.paper_style_var.get(),
                        selected_font=font_path,
                        paper_size=self.paper_size_var.get(),
                        custom_width=self.custom_width_var.get(),
                        custom_height=self.custom_height_var.get(),
                        black_white=True
                    )
                else:
                    export_image = self.current_image
                
                # Tạo document Word
                doc = Document()
                
                # Lưu image tạm
                temp_img_path = "temp_handwriting.png"
                export_image.save(temp_img_path)
                
                # Thêm image vào Word
                paragraph = doc.add_paragraph()
                run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
                run.add_picture(temp_img_path, width=Inches(6))
                
                # Lưu document
                doc.save(file_path)
                
                # Xóa file tạm
                if os.path.exists(temp_img_path):
                    os.remove(temp_img_path)
                
                messagebox.showinfo("Thành công", f"Đã lưu: {os.path.basename(file_path)}")
                self.status_bar.config(text=f"Đã xuất Word: {os.path.basename(file_path)}")
            except Exception as e:
                messagebox.showerror("Lỗi", f"Không thể tạo Word: {e}")
                
    def export_svg(self):
        """Xuất file SVG"""
        text = self.text_area.get('1.0', tk.END).strip()
        if not text:
            messagebox.showwarning("Cảnh báo", "Vui lòng nhập văn bản!")
            return
            
        file_path = filedialog.asksaveasfilename(
            title="Lưu file SVG",
            defaultextension=".svg",
            filetypes=[("SVG files", "*.svg"), ("All files", "*.*")]
        )
        if file_path:
            try:
                # Lấy thông tin font
                selected_display_name = self.selected_font.get()
                font_path, font_lang = None, None
                if hasattr(self, 'font_lang_map') and selected_display_name in self.font_lang_map:
                    font_path, font_lang = self.font_lang_map[selected_display_name]
                else:
                    font_path = None
                    font_lang = self.font_lang
                
                # Tạo SVG content
                svg_content = self.generate_svg_content(
                    text=text,
                    font_size=self.font_size_var.get(),
                    line_spacing=self.line_spacing_var.get(),
                    margin=self.margin_var.get(),
                    ink_color=self.ink_color,
                    paper_style=self.paper_style_var.get(),
                    selected_font=font_path,
                    paper_size=self.paper_size_var.get(),
                    custom_width=self.custom_width_var.get(),
                    custom_height=self.custom_height_var.get(),
                    black_white=self.black_white_var.get()
                )
                
                # Lưu SVG vào file
                with open(file_path, 'w', encoding='utf-8') as f:
                    f.write(svg_content)
                
                messagebox.showinfo("Thành công", f"Đã lưu: {os.path.basename(file_path)}")
                self.status_bar.config(text=f"Đã xuất SVG: {os.path.basename(file_path)}")
            except Exception as e:
                messagebox.showerror("Lỗi", f"Không thể tạo SVG: {e}")
    
    def generate_svg_content(self, text, font_size, line_spacing, margin, ink_color, 
                           paper_style, selected_font, paper_size, custom_width, 
                           custom_height, black_white):
        """Tạo nội dung SVG"""
        # Tính toán kích thước canvas
        line_height = int(font_size * line_spacing)
        canvas_width, canvas_height = self.get_canvas_size(paper_size, custom_width, custom_height, 1, line_height, margin)
        
        # Tính toán số ký tự tối đa trên mỗi dòng
        available_width = canvas_width - margin * 2 - 60
        estimated_char_width = font_size * 0.6
        max_chars_per_line = int(available_width / estimated_char_width)
        
        # Xử lý text
        wrapper = textwrap.TextWrapper(width=max_chars_per_line)
        lines = []
        for paragraph in text.split('\n'):
            if paragraph.strip():
                lines.extend(wrapper.wrap(paragraph))
            else:
                lines.append('')
        
        # Tính lại chiều cao canvas
        canvas_width, canvas_height = self.get_canvas_size(paper_size, custom_width, custom_height, len(lines), line_height, margin)
        
        # Tạo SVG header
        svg_content = f'''<?xml version="1.0" encoding="UTF-8"?>
<svg width="{canvas_width}" height="{canvas_height}" xmlns="http://www.w3.org/2000/svg">
  <defs>
    <style>
      .text {{ font-family: Arial, sans-serif; }}
      .paper-bg {{ fill: white; }}
    </style>
  </defs>
  
  <!-- Background -->
  <rect width="{canvas_width}" height="{canvas_height}" class="paper-bg"/>
'''
        
        # Thêm background pattern cho paper style
        if paper_style == "lined":
            # Vẽ dòng kẻ ngang
            for i in range(margin + line_height, canvas_height - margin, line_height):
                svg_content += f'  <line x1="{margin}" y1="{i}" x2="{canvas_width - margin}" y2="{i}" stroke="#c8c8ff" stroke-width="1"/>\n'
            # Vẽ lề đỏ
            svg_content += f'  <line x1="{margin + 50}" y1="{margin}" x2="{margin + 50}" y2="{canvas_height - margin}" stroke="#ff6464" stroke-width="2"/>\n'
        elif paper_style == "olined":
            # Vẽ giấy 4 ô ly
            for i in range(margin, canvas_height - margin, line_height // 4):
                color = "#6464c8" if (i - margin) % line_height == 0 else "#b4b4ff"
                width = "2" if (i - margin) % line_height == 0 else "1"
                svg_content += f'  <line x1="{margin}" y1="{i}" x2="{canvas_width - margin}" y2="{i}" stroke="{color}" stroke-width="{width}"/>\n'
            # Vẽ dòng dọc
            for x in range(margin + 60, canvas_width - margin, 80):
                color = "#9696c8" if (x - margin - 60) % 160 == 0 else "#c8c8dc"
                width = "2" if (x - margin - 60) % 160 == 0 else "1"
                svg_content += f'  <line x1="{x}" y1="{margin}" x2="{x}" y2="{canvas_height - margin}" stroke="{color}" stroke-width="{width}"/>\n'
            # Lề đỏ
            svg_content += f'  <line x1="{margin + 50}" y1="{margin}" x2="{margin + 50}" y2="{canvas_height - margin}" stroke="#ff6464" stroke-width="2"/>\n'
        
        # Thêm text
        y_position = margin + line_height
        for line in lines:
            if line.strip():
                x_position = margin + 60
                # Tạo text với biến thể nhẹ
                for char in line:
                    char_x = x_position + random.randint(-2, 2)
                    char_y = y_position + random.randint(-2, 2)
                    
                    # Màu text
                    if black_white:
                        text_color = "#000000"
                    else:
                        ink_variation = random.randint(-20, 20)
                        r = max(0, min(255, ink_color[0] + ink_variation))
                        g = max(0, min(255, ink_color[1] + ink_variation))
                        b = max(0, min(255, ink_color[2] + ink_variation))
                        text_color = f"rgb({r},{g},{b})"
                    
                    # Kích thước font biến thể
                    size_variation = random.randint(-2, 2)
                    char_font_size = max(8, font_size + size_variation)
                    
                    svg_content += f'  <text x="{char_x}" y="{char_y}" font-size="{char_font_size}" fill="{text_color}" class="text">{char}</text>\n'
                    
                    # Tính khoảng cách đến ký tự tiếp theo
                    char_width = char_font_size * 0.6
                    spacing_variation = random.randint(-2, 3)
                    x_position += char_width + spacing_variation
            
            y_position += line_height
        
        svg_content += '</svg>'
        return svg_content
    
    def save_settings(self):
        """Lưu cài đặt"""
        selected_display_name = self.selected_font.get()
        font_path, font_lang = None, None
        if hasattr(self, 'font_lang_map') and selected_display_name in self.font_lang_map:
            font_path, font_lang = self.font_lang_map[selected_display_name]
        
        settings = {
            'font_size': self.font_size_var.get(),
            'line_spacing': self.line_spacing_var.get(),
            'margin': self.margin_var.get(),
            'ink_color': self.ink_color,
            'paper_style': self.paper_style_var.get(),
            'selected_font': font_path,
            'paper_size': self.paper_size_var.get(),
            'custom_width': self.custom_width_var.get(),
            'custom_height': self.custom_height_var.get(),
            'black_white': self.black_white_var.get(),
            'dark_mode': self.dark_mode_var.get()
        }
        
        try:
            with open('handwriting_settings.json', 'w', encoding='utf-8') as f:
                json.dump(settings, f, indent=2, ensure_ascii=False)
            messagebox.showinfo("Thành công", "Đã lưu cài đặt!")
            self.status_bar.config(text="Đã lưu cài đặt")
        except Exception as e:
            messagebox.showerror("Lỗi", f"Không thể lưu cài đặt: {e}")
            
    def load_settings(self):
        """Tải cài đặt"""
        try:
            if os.path.exists('handwriting_settings.json'):
                with open('handwriting_settings.json', 'r', encoding='utf-8') as f:
                    settings = json.load(f)
                # Bổ sung các trường mới nếu thiếu
                if 'paper_size' not in settings:
                    settings['paper_size'] = 'A4'
                if 'custom_width' not in settings:
                    settings['custom_width'] = 800
                if 'custom_height' not in settings:
                    settings['custom_height'] = 600
                if 'black_white' not in settings:
                    settings['black_white'] = False
                if 'dark_mode' not in settings:
                    settings['dark_mode'] = False
                return settings
        except Exception as e:
            print(f"Lỗi tải cài đặt: {e}")
        # Default settings
        return {
            'font_size': 28,
            'line_spacing': 1.5,
            'margin': 50,
            'ink_color': (25, 25, 112),
            'paper_style': 'lined',
            'selected_font': None,
            'paper_size': 'A4',
            'custom_width': 800,
            'custom_height': 600,
            'black_white': False,
            'dark_mode': False
        }
        
    def load_settings_to_ui(self):
        """Áp dụng cài đặt lên UI"""
        self.font_size_var.set(self.settings['font_size'])
        self.line_spacing_var.set(self.settings['line_spacing'])
        self.margin_var.set(self.settings['margin'])
        self.ink_color = tuple(self.settings['ink_color'])
        self.paper_style_var.set(self.settings['paper_style'])
        self.paper_size_var.set(self.settings['paper_size'])
        self.custom_width_var.set(self.settings['custom_width'])
        self.custom_height_var.set(self.settings['custom_height'])
        
        # Áp dụng tùy chọn đen trắng
        if 'black_white' in self.settings:
            self.black_white_var.set(self.settings['black_white'])
        else:
            self.black_white_var.set(False)
        
        # Áp dụng tùy chọn dark mode
        if 'dark_mode' in self.settings:
            self.dark_mode_var.set(self.settings['dark_mode'])
            if self.settings['dark_mode']:
                self.toggle_dark_mode()  # Áp dụng dark mode nếu được lưu
        else:
            self.dark_mode_var.set(False)
        
        # Cập nhật màu button
        self.color_button.config(bg=self.rgb_to_hex(self.ink_color))
        
        # Cập nhật labels
        self.font_size_label.config(text=str(self.settings['font_size']))
        self.line_spacing_label.config(text=f"{self.settings['line_spacing']:.1f}")
        self.margin_label.config(text=str(self.settings['margin']))
        
        # Cập nhật font được chọn
        saved_font_path = self.settings.get('selected_font')
        if saved_font_path and saved_font_path in self.generator.fonts:
            try:
                font_index = self.generator.fonts.index(saved_font_path)
                if self.font_combo['values'] and font_index < len(self.font_combo['values']):
                    self.font_combo.current(font_index)
            except ValueError:
                if self.font_combo['values']:
                    self.font_combo.current(0)
        elif self.font_combo['values']:
            self.font_combo.current(0)
        # Nếu không có font nào, không gọi current
        self.update_font_demo()

    def setup_font_list(self):
        """Luôn hiển thị tất cả font của mọi ngôn ngữ, có nhãn nước"""
        font_display_names = []
        font_lang_map = {}
        for lang, _ in FONT_LANGS:
            temp_gen = HandwritingGenerator(font_lang=lang)
            for p in temp_gen.fonts:
                display = temp_gen.font_names.get(p)
                if display:
                    label = LANG_LABELS.get(lang, lang)
                    display_name = f"[{label}] {display}"
                    font_display_names.append(display_name)
                    font_lang_map[display_name] = (p, lang)
        self.font_combo['values'] = font_display_names
        self.font_lang_map = font_lang_map
        num_fonts = len(font_display_names)
        self.font_count_label.config(text=f"Có {num_fonts} font khả dụng" if num_fonts else "Không có font nào")
        if num_fonts == 0:
            self.font_hint_label.config(text="Chưa có font hợp lệ. Kéo-thả hoặc copy file .ttf/.otf vào thư mục fonts/[ngôn ngữ]/ để sử dụng.")
        else:
            self.font_hint_label.config(text="")
        if font_display_names:
            self.font_combo.current(len(font_display_names)-1)
            self.update_font_file_label()
        else:
            self.font_file_label.config(text="")
        self.font_combo.bind('<<ComboboxSelected>>', lambda e: self.update_font_file_label())
        self.update_font_demo()
    
    def update_font_demo(self, event=None):
        """Cập nhật demo font"""
        try:
            selected_display_name = self.selected_font.get()
            font_path, font_lang = None, None
            if hasattr(self, 'font_lang_map') and selected_display_name in self.font_lang_map:
                font_path, font_lang = self.font_lang_map[selected_display_name]
            demo_text = self.demo_text_var.get() if hasattr(self, 'demo_text_var') else "Xin chào! Hello! 123"
            if font_path:
                try:
                    # Tạo generator tạm đúng ngôn ngữ để load font demo
                    temp_gen = HandwritingGenerator(font_lang=font_lang)
                    demo_img = Image.new('RGB', (280, 50), color=(255, 255, 255))
                    demo_draw = ImageDraw.Draw(demo_img)
                    demo_font = ImageFont.truetype(font_path, 16, encoding='unic')
                    demo_draw.text((10, 10), demo_text, font=demo_font, fill=(0, 0, 0))
                    demo_photo = ImageTk.PhotoImage(demo_img)
                    self.font_demo_label.config(image=demo_photo, text="")
                    self.font_demo_label.image = demo_photo
                except Exception as e:
                    print(f"Lỗi tạo demo font: {e}")
                    self.font_demo_label.config(image="", text="Font không khả dụng", font=('Arial', 10), foreground='red')
            else:
                self.font_demo_label.config(image="", text="Font mặc định", font=('Arial', 10), foreground='blue')
        except Exception as e:
            print(f"Lỗi cập nhật font demo: {e}")
            self.font_demo_label.config(image="", text="Lỗi font", font=('Arial', 10), foreground='red')

    def on_canvas_press(self, event):
        """Ghi nhận vị trí bắt đầu khi nhấn chuột để di chuyển."""
        self.preview_canvas.scan_mark(event.x, event.y)

    def on_canvas_drag(self, event):
        """Di chuyển ảnh trên canvas khi kéo chuột."""
        self.preview_canvas.scan_dragto(event.x, event.y, gain=1)

    def on_canvas_enter(self, event):
        """Đổi con trỏ chuột khi di vào vùng xem trước để báo hiệu có thể di chuyển."""
        self.preview_canvas.config(cursor="fleur")

    def on_canvas_leave(self, event):
        """Khôi phục con trỏ chuột khi rời khỏi vùng xem trước."""
        self.preview_canvas.config(cursor="")

    def on_drop_font_file(self, event):
        """Xử lý khi người dùng kéo-thả file font vào cửa sổ"""
        import shutil
        files = self.root.tk.splitlist(event.data)
        font_added = False
        last_font_name = None
        for file_path in files:
            if file_path.lower().endswith(('.ttf', '.otf')):
                try:
                    dest_dir = os.path.join('fonts', self.font_lang)
                    if not os.path.exists(dest_dir):
                        os.makedirs(dest_dir)
                    dest = os.path.join(dest_dir, os.path.basename(file_path))
                    shutil.copy(file_path, dest)
                    font_added = True
                    last_font_name = os.path.splitext(os.path.basename(file_path))[0]
                except Exception as e:
                    messagebox.showerror("Lỗi", f"Không thể thêm font: {e}")
        if font_added:
            messagebox.showinfo("Thành công", "Đã thêm font mới! Danh sách font sẽ được cập nhật.")
            self.generator.load_fonts(font_lang=self.font_lang)
            self.setup_font_list()
            # Tự động chọn font vừa thêm
            if last_font_name:
                font_display_names = [self.generator.font_names.get(p) for p in self.generator.fonts]
                if last_font_name in font_display_names:
                    self.font_combo.current(font_display_names.index(last_font_name))
                    self.update_font_file_label()
        else:
            messagebox.showwarning("Cảnh báo", "Chỉ hỗ trợ file .ttf hoặc .otf khi kéo-thả vào.")

    def on_paper_size_change(self, event=None):
        if self.paper_size_var.get() == "Custom":
            self.custom_size_frame.grid()
        else:
            self.custom_size_frame.grid_remove()

    def update_font_file_label(self):
        # Hiển thị tên file font khi chọn
        idx = self.font_combo.current()
        if idx >= 0:
            font_display_names = [self.generator.font_names.get(p) for p in self.generator.fonts]
            font_paths = list(self.generator.font_names.keys())
            if idx < len(font_paths):
                font_path = font_paths[idx]
                if font_path:
                    self.font_file_label.config(text=f"File: {os.path.basename(font_path)}")
                else:
                    self.font_file_label.config(text="Font mặc định (Lỗi)")
        else:
            self.font_file_label.config(text="")

    def update_ui_language(self):
        """Cập nhật ngôn ngữ UI"""
        self.L = LANGS[self.lang]
        self.root.title(self.L['title'])
        # Cập nhật các label khác nếu cần

    # Template System Functions
    def save_template(self):
        """Lưu template hiện tại"""
        template_name = self.template_name_var.get().strip()
        if not template_name:
            messagebox.showwarning("Cảnh báo", "Vui lòng nhập tên template!")
            return
        
        # Lấy thông tin font hiện tại
        selected_display_name = self.selected_font.get()
        font_path, font_lang = None, None
        if hasattr(self, 'font_lang_map') and selected_display_name in self.font_lang_map:
            font_path, font_lang = self.font_lang_map[selected_display_name]
        
        template_data = {
            'name': template_name,
            'font_size': self.font_size_var.get(),
            'line_spacing': self.line_spacing_var.get(),
            'margin': self.margin_var.get(),
            'ink_color': self.ink_color,
            'paper_style': self.paper_style_var.get(),
            'selected_font': font_path,
            'font_lang': font_lang,
            'paper_size': self.paper_size_var.get(),
            'custom_width': self.custom_width_var.get(),
            'custom_height': self.custom_height_var.get(),
            'black_white': self.black_white_var.get(),
            'created_date': str(datetime.datetime.now())
        }
        
        try:
            # Tải danh sách template hiện tại
            templates = self.load_templates_list()
            templates[template_name] = template_data
            
            # Lưu vào file
            with open('templates.json', 'w', encoding='utf-8') as f:
                json.dump(templates, f, indent=2, ensure_ascii=False)
            
            messagebox.showinfo("Thành công", f"Đã lưu template '{template_name}'!")
            self.refresh_template_list()
            self.template_name_var.set("")  # Xóa tên template
            self.status_bar.config(text=f"Đã lưu template: {template_name}")
        except Exception as e:
            messagebox.showerror("Lỗi", f"Không thể lưu template: {e}")
    
    def load_template(self):
        """Tải template từ file"""
        file_path = filedialog.askopenfilename(
            title="Chọn file template",
            filetypes=[("JSON files", "*.json"), ("All files", "*.*")]
        )
        if file_path:
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    template_data = json.load(f)
                self.apply_template(template_data)
                messagebox.showinfo("Thành công", "Đã tải template!")
            except Exception as e:
                messagebox.showerror("Lỗi", f"Không thể tải template: {e}")
    
    def load_template_from_list(self, event=None):
        """Tải template từ danh sách"""
        selection = self.template_listbox.curselection()
        if not selection:
            return
        
        template_name = self.template_listbox.get(selection[0])
        templates = self.load_templates_list()
        
        if template_name in templates:
            self.apply_template(templates[template_name])
            messagebox.showinfo("Thành công", f"Đã tải template '{template_name}'!")
        else:
            messagebox.showerror("Lỗi", "Template không tồn tại!")
    
    def delete_template(self):
        """Xóa template"""
        selection = self.template_listbox.curselection()
        if not selection:
            messagebox.showwarning("Cảnh báo", "Vui lòng chọn template để xóa!")
            return
        
        template_name = self.template_listbox.get(selection[0])
        result = messagebox.askyesno("Xác nhận", f"Bạn có chắc muốn xóa template '{template_name}'?")
        
        if result:
            try:
                templates = self.load_templates_list()
                if template_name in templates:
                    del templates[template_name]
                    
                    with open('templates.json', 'w', encoding='utf-8') as f:
                        json.dump(templates, f, indent=2, ensure_ascii=False)
                    
                    messagebox.showinfo("Thành công", f"Đã xóa template '{template_name}'!")
                    self.refresh_template_list()
                else:
                    messagebox.showerror("Lỗi", "Template không tồn tại!")
            except Exception as e:
                messagebox.showerror("Lỗi", f"Không thể xóa template: {e}")
    
    def load_templates_list(self):
        """Tải danh sách template từ file"""
        try:
            if os.path.exists('templates.json'):
                with open('templates.json', 'r', encoding='utf-8') as f:
                    return json.load(f)
        except Exception as e:
            print(f"Lỗi tải templates: {e}")
        return {}
    
    def refresh_template_list(self):
        """Làm mới danh sách template"""
        self.template_listbox.delete(0, tk.END)
        templates = self.load_templates_list()
        
        for template_name in sorted(templates.keys()):
            template_data = templates[template_name]
            created_date = template_data.get('created_date', 'Unknown')
            display_text = f"{template_name} ({created_date[:10]})"
            self.template_listbox.insert(tk.END, display_text)
    
    def apply_template(self, template_data):
        """Áp dụng template lên UI"""
        try:
            self.font_size_var.set(template_data.get('font_size', 28))
            self.line_spacing_var.set(template_data.get('line_spacing', 1.5))
            self.margin_var.set(template_data.get('margin', 50))
            self.ink_color = tuple(template_data.get('ink_color', (25, 25, 112)))
            self.paper_style_var.set(template_data.get('paper_style', 'lined'))
            self.paper_size_var.set(template_data.get('paper_size', 'A4'))
            self.custom_width_var.set(template_data.get('custom_width', 800))
            self.custom_height_var.set(template_data.get('custom_height', 600))
            self.black_white_var.set(template_data.get('black_white', False))
            
            # Cập nhật màu button
            self.color_button.config(bg=self.rgb_to_hex(self.ink_color))
            
            # Cập nhật labels
            self.font_size_label.config(text=str(template_data.get('font_size', 28)))
            self.line_spacing_label.config(text=f"{template_data.get('line_spacing', 1.5):.1f}")
            self.margin_label.config(text=str(template_data.get('margin', 50)))
            
            # Cập nhật font được chọn
            saved_font_path = template_data.get('selected_font')
            if saved_font_path and hasattr(self, 'font_lang_map'):
                # Tìm font trong danh sách
                for display_name, (font_path, font_lang) in self.font_lang_map.items():
                    if font_path == saved_font_path:
                        self.font_combo.set(display_name)
                        break
                else:
                    if self.font_combo['values']:
                        self.font_combo.current(0)
            
            self.update_font_demo()
            self.status_bar.config(text=f"Đã áp dụng template: {template_data.get('name', 'Unknown')}")
            
        except Exception as e:
            messagebox.showerror("Lỗi", f"Không thể áp dụng template: {e}")

    def toggle_dark_mode(self):
        """Chuyển đổi giữa light mode và dark mode (cải tiến toàn diện)"""
        import tkinter.ttk as ttk
        style = ttk.Style()
        if self.dark_mode_var.get():
            # Định nghĩa style dark cho ttk
            style.theme_use('clam')
            style.configure('TFrame', background='#232629')
            style.configure('TLabelframe', background='#232629', foreground='#fff')
            style.configure('TLabel', background='#232629', foreground='#fff')
            style.configure('TButton', background='#404040', foreground='#fff')
            style.configure('TEntry', fieldbackground='#404040', foreground='#fff')
            style.configure('TCombobox', fieldbackground='#404040', foreground='#fff', background='#404040')
            style.configure('TCheckbutton', background='#232629', foreground='#fff')
            style.configure('TProgressbar', background='#404040', troughcolor='#232629')
            style.configure('TScale', background='#232629', troughcolor='#404040')
            style.configure('Horizontal.TScrollbar', background='#404040', troughcolor='#232629')
            style.configure('Vertical.TScrollbar', background='#404040', troughcolor='#232629')
            bg_color = '#232629'
            fg_color = '#ffffff'
            entry_bg = '#404040'
            entry_fg = '#ffffff'
            canvas_bg = '#181a1b'
        else:
            style.theme_use('default')
            style.configure('TFrame', background='#f0f0f0')
            style.configure('TLabelframe', background='#f0f0f0', foreground='#000')
            style.configure('TLabel', background='#f0f0f0', foreground='#000')
            style.configure('TButton', background='#e0e0e0', foreground='#000')
            style.configure('TEntry', fieldbackground='#ffffff', foreground='#000')
            style.configure('TCombobox', fieldbackground='#ffffff', foreground='#000', background='#ffffff')
            style.configure('TCheckbutton', background='#f0f0f0', foreground='#000')
            style.configure('TProgressbar', background='#e0e0e0', troughcolor='#f0f0f0')
            style.configure('TScale', background='#f0f0f0', troughcolor='#e0e0e0')
            style.configure('Horizontal.TScrollbar', background='#e0e0e0', troughcolor='#f0f0f0')
            style.configure('Vertical.TScrollbar', background='#e0e0e0', troughcolor='#f0f0f0')
            bg_color = '#f0f0f0'
            fg_color = '#000000'
            entry_bg = '#ffffff'
            entry_fg = '#000000'
            canvas_bg = '#ffffff'
        # Đệ quy set màu cho toàn bộ widget
        self.apply_dark_mode_to_widget(self.root, bg_color, fg_color, entry_bg, entry_fg, canvas_bg)
        # Status bar
        self.status_bar.config(background=bg_color, foreground=fg_color)
        mode_text = "Dark Mode" if self.dark_mode_var.get() else "Light Mode"
        self.status_bar.config(text=f"Đã chuyển sang {mode_text}")

    def apply_dark_mode_to_widget(self, widget, bg, fg, entry_bg, entry_fg, canvas_bg):
        try:
            widget_type = widget.winfo_class()
            # Các widget thường
            if widget_type in ['Frame', 'Labelframe']:
                widget.configure(bg=bg)
            elif widget_type in ['Label']:
                widget.configure(bg=bg, fg=fg)
            elif widget_type in ['Button']:
                widget.configure(bg=bg, fg=fg, activebackground=entry_bg, activeforeground=fg)
            elif widget_type in ['Entry']:
                widget.configure(bg=entry_bg, fg=entry_fg, insertbackground=fg)
            elif widget_type in ['Text']:
                widget.configure(bg=entry_bg, fg=entry_fg, insertbackground=fg, selectbackground=bg, selectforeground=fg)
            elif widget_type in ['Canvas']:
                widget.configure(bg=canvas_bg)
            elif widget_type in ['Listbox']:
                widget.configure(bg=entry_bg, fg=entry_fg, selectbackground=bg, selectforeground=fg)
            elif widget_type in ['Scrollbar']:
                widget.configure(bg=bg, troughcolor=entry_bg)
            # Đệ quy cho widget con
            for child in widget.winfo_children():
                self.apply_dark_mode_to_widget(child, bg, fg, entry_bg, entry_fg, canvas_bg)
        except Exception:
            pass

    def get_canvas_size(self, paper_size, custom_width, custom_height, num_lines, line_height, margin):
        """Tính toán kích thước canvas (copy từ HandwritingGenerator)"""
        if paper_size == "A4":
            # 794x1123 px ~ 210x297mm ở 96dpi
            return 794, max(1123, num_lines * line_height + margin * 2)
        elif paper_size == "A5":
            # 559x794 px ~ 148x210mm ở 96dpi
            return 559, max(794, num_lines * line_height + margin * 2)
        elif paper_size == "Custom":
            return custom_width, custom_height
        else:
            return 800, max(600, num_lines * line_height + margin * 2)

def main():
    """Chạy ứng dụng"""
    try:
        root = tk.Tk()
        app = HandwritingGUI(root)
        root.mainloop()
    except Exception as e:
        print(f"Lỗi chạy ứng dụng: {e}")
        messagebox.showerror("Lỗi", f"Không thể chạy ứng dụng: {e}")

if __name__ == "__main__":
    main()